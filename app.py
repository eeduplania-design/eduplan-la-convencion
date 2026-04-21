import streamlit as st
from zhipuai import ZhipuAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import io
import re
import datetime
import requests
import urllib.parse
import time

# --- 1. CONFIGURACIÓN INICIAL Y DATOS MAESTROS ---
st.set_page_config(page_title="EDUPLAN IA - MINEDU", layout="wide", page_icon="🇵🇪")

NOMBRE_APP = "EDUPLAN IA - SISTEMA EXPERTO MINEDU"
ANIO_ACTUAL = datetime.datetime.now().year

ENFOQUES_TRANSVERSALES = [
    "De derechos", "Inclusivo o de Atención a la diversidad", 
    "Intercultural", "Igualdad de género", "Ambiental", 
    "Orientación al bien común", "Búsqueda de la Excelencia"
]

INSTRUMENTOS_EVALUACION = [
    "Rúbricas",
    "Listas de cotejo",
    "Escalas de valoración",
    "Fichas de observación",
    "Portafolio"
]

NIVELES_GRADOS = {
    "Inicial": ["3 años", "4 años", "5 años"],
    "Primaria": ["1ro", "2do", "3ro", "4to", "5to", "6to"],
    "Secundaria": ["1ro", "2do", "3ro", "4to", "5to"]
}

AREAS_NIVEL = {
    "Inicial": ["Personal Social", "Psicomotriz", "Comunicación", "Castellano como segunda lengua", "Descubrimiento del Mundo", "Matemática", "Ciencia y Tecnología"],
    "Primaria": ["Matemática", "Comunicación", "Inglés", "Personal Social", "Arte y Cultura", "Ciencia y Tecnología", "Educación Física", "Educación Religiosa", "Tutoría"],
    "Secundaria": ["Matemática", "Comunicación", "Inglés", "Arte y Cultura", "Ciencias Sociales", "Desarrollo Personal, Ciudadanía y Cívica (DPCC)", "Educación Física", "Educación Religiosa", "Ciencia y Tecnología", "Educación para el Trabajo", "Tutoría"]
}

# --- 2. LÓGICA PEDAGÓGICA (CNEB) ---
def obtener_ciclo(nivel, grado):
    if nivel == "Inicial": return "Ciclo II"
    if nivel == "Primaria":
        if grado in ["1ro", "2do"]: return "Ciclo III"
        if grado in ["3ro", "4to"]: return "Ciclo IV"
        if grado in ["5to", "6to"]: return "Ciclo V"
    if nivel == "Secundaria":
        if grado in ["1ro", "2do"]: return "Ciclo VI"
        if grado in ["3ro", "4to", "5to"]: return "Ciclo VII"
    return "Ciclo no definido"

def generar_prompt_cneb(tipo_doc, datos):
    nivel = datos.get("nivel")
    grado = datos.get("grado")
    area = datos.get("area")
    ciclo = obtener_ciclo(nivel, grado)
    instrumento = datos.get("instrumento", "Rúbricas")
    
    base_prompt = f"""Eres un Especialista Senior del MINEDU de Perú con más de 20 años de experiencia. 
Tu tarea es redactar un/una '{tipo_doc}' de nivel profesional, alineado al CNEB (Currículo Nacional).
Estrictamente adaptado a:
- Nivel: {nivel}
- Grado: {grado} (Ciclo: {ciclo})
- Área: {area}
- Tema/Título: {datos.get('tema', 'No especificado')}

REGLAS DE FORMATO (CRÍTICAS PARA EL PARSEO A WORD):
1. Usa Markdown estándar. Usa # para Título Principal, ## para Secciones, ### para Subsecciones.
2. Usa tablas Markdown para matrices, competencias, criterios y rúbricas. (Ejemplo: | Col1 | Col2 |). NUNCA dejes celdas vacías, pon "-".
3. Para Sesiones de Aprendizaje, debes incluir detalladamente: Inicio, Desarrollo y Cierre.
4. Al final de la Sesión, debes generar DOS ANEXOS:
   - ANEXO 1: Instrumento de evaluación ({instrumento} en formato TABLA). Aplica el formato correspondiente.
   - ANEXO 2: Ficha de Aplicación para el estudiante adaptada al nivel cognitivo de {grado} de {nivel}.
5. Si necesitas que en la Ficha haya una imagen ilustrativa, inserta EXACTAMENTE esta etiqueta: [IMAGEN_SUGERIDA: breve descripción en ingles].

Información ingresada por el docente:
- Docente: {datos.get('docente')}
- I.E.: {datos.get('ie')}
- Enfoque Transversal: {datos.get('enfoque')}
- Situación Significativa / Contexto: {datos.get('contexto', 'Contexto local estándar')}
- Duración: {datos.get('duracion', '2 horas pedagógicas')}
"""

    if tipo_doc == "Programación Anual":
        base_prompt += "\nEstructura: 1. Datos Generales. 2. Descripción General. 3. Organización de Unidades/Experiencias. 4. Matriz de Competencias y Enfoques. 5. Evaluación. 6. Materiales."
    elif tipo_doc == "Unidad Didáctica":
        base_prompt += f"\nEstructura: 1. Datos Generales. 2. Situación Significativa. 3. Propósitos de Aprendizaje. 4. Criterios, Evidencias e Instrumentos. 5. Secuencia de Sesiones (Tabla con N° de sesión, Título y Descripción breve). Producto final: {datos.get('producto', 'No especificado')}."
    elif tipo_doc == "Sesión de Aprendizaje":
        base_prompt += f"""
Estructura: 1. Datos Informativos. 2. Título. 3. Propósitos (Tabla Competencia/Capacidad/Desempeño/Criterio/Evidencia). 4. Enfoques. 5. Preparación. 6. Momentos (Inicio, Desarrollo -con procesos didácticos del área-, Cierre). 7. Reflexiones. 8. Anexos (Instrumento y Ficha).

✅ RECOMENDACIONES DEL MINEDU PARA EL INSTRUMENTO SELECCIONADO ({instrumento}):
- Los instrumentos deben estar alineados a las competencias del CNEB.
- Se prioriza la evaluación formativa: retroalimentación constante y oportuna.
- Si elegiste Rúbricas: Construye una Matriz con criterios y niveles de logro (inicio, en proceso, logrado, destacado).
- Si elegiste Listas de cotejo: Crea un registro de verificación de aspectos cumplidos o no (sí/no) con indicadores claros.
- Si elegiste Escalas de valoración: Mide frecuencia o calidad en una escala (siempre, a veces, nunca).
- Si elegiste Fichas de observación: Haz un documento para registrar conductas o desempeños observados.
- Si elegiste Portafolio: Lista la colección de evidencias del aprendizaje sugeridas.
- Aplica este diseño detallado al construir el ANEXO 1.
"""

    return base_prompt

# --- 3. GENERACIÓN DE WORD (PARSER CALIDAD IMPRENTA) ---
def add_header_footer(doc):
    """Añade encabezado y pie de página institucional"""
    section = doc.sections[0]
    
    # Márgenes (A4 Normal)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # Encabezado
    header = section.header
    h_p = header.paragraphs[0]
    h_p.text = "MINISTERIO DE EDUCACIÓN DEL PERÚ\nPlanificación Curricular - CNEB"
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in h_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

def insertar_imagen_pollinations(doc, prompt_imagen):
    """Busca e inserta imagen desde Pollinations AI con validación de seguridad"""
    try:
        query = urllib.parse.quote(f"{prompt_imagen}, educational vector style, flat illustration, for kids, white background")
        url = f"https://image.pollinations.ai/prompt/{query}?width=500&height=350&nologo=true"
        response = requests.get(url, timeout=10) # Aumentamos un poco el timeout
        
        # Validar que la respuesta sea un 200 OK y que el contenido sea realmente una imagen
        if response.status_code == 200 and 'image' in response.headers.get('Content-Type', ''):
            image_stream = io.BytesIO(response.content)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(4.5))
        else:
            doc.add_paragraph("[Imagen no disponible temporalmente]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        # Falla silenciosa controlada, inserta un recuadro de texto
        doc.add_paragraph("[Espacio para imagen referencial]").alignment = WD_ALIGN_PARAGRAPH.CENTER

def markdown_to_docx(markdown_text):
    """Convierte Markdown a un DOCX estructurado y profesional."""
    doc = Document()
    add_header_footer(doc)
    
    lines = markdown_text.split('\n')
    in_table = False
    table_data = []

    def procesar_tabla(datos):
        if not datos: return
        try:
            # Filtrar la fila de separadores de markdown (|---|---|)
            datos = [row for row in datos if not re.match(r'^\|[\-\|\s]+\|$', row.strip())]
            if not datos: return

            # Contar columnas usando la primera fila válida
            headers = [c.strip() for c in datos[0].split('|')[1:-1]]
            num_cols = len(headers)
            
            if num_cols == 0: return # Evitar procesar si no hay columnas válidas
            
            table = doc.add_table(rows=len(datos), cols=num_cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(datos):
                cells_text = [c.strip() for c in row.split('|')[1:-1]]
                for j in range(min(num_cols, len(cells_text))):
                    cell = table.cell(i, j)
                    text = cells_text[j].replace('**', '')
                    cell.text = text
                    
                    # Formato Encabezado de Tabla (Fila 0)
                    if i == 0:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="D9E2F3"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                    else:
                        # Contenido normal
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            doc.add_paragraph()
        except Exception as e:
            # SALVAVIDAS: Si la tabla viene mal formada por la IA, se imprime como texto en lugar de romper la app.
            for row in datos:
                doc.add_paragraph(row.replace('|', ' ')).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()

    for line in lines:
        line_stripped = line.strip()
        
        # 1. Manejo de Tablas
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            in_table = True
            table_data.append(line_stripped)
            continue
        elif in_table:
            procesar_tabla(table_data)
            in_table = False
            table_data = []
            if not line_stripped: continue

        # 2. Manejo de Imágenes
        img_match = re.search(r'\[IMAGEN_SUGERIDA:\s*(.*?)\]', line_stripped)
        if img_match:
            insertar_imagen_pollinations(doc, img_match.group(1))
            continue

        # 3. Encabezados (Headings)
        if line_stripped.startswith('### '):
            p = doc.add_heading(line_stripped.replace('### ', ''), level=3)
        elif line_stripped.startswith('## '):
            p = doc.add_heading(line_stripped.replace('## ', ''), level=2)
            for run in p.runs: run.font.color.rgb = RGBColor(0, 51, 102) # Azul MINEDU
        elif line_stripped.startswith('# '):
            p = doc.add_heading(line_stripped.replace('# ', ''), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: 
                run.font.color.rgb = RGBColor(200, 16, 46) # Rojo MINEDU
        
        # 4. Párrafos y Listas
        else:
            if line_stripped == "": continue
            
            if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = line_stripped[2:]
            elif re.match(r'^\d+\.', line_stripped):
                p = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                text = line_stripped

            # Procesar Negritas en línea (**texto**)
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Si el documento termina en tabla
    if in_table:
        procesar_tabla(table_data)

    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. INTERFAZ DE USUARIO (STREAMLIT) ---
st.markdown("""
    <style>
    .header-box {
        background: linear-gradient(135deg, #003366 0%, #1e40af 100%);
        padding: 25px; border-radius: 10px; color: white; text-align: center;
        margin-bottom: 25px; border-bottom: 5px solid #C8102E;
    }
    .header-box h1 { font-size: 2.2rem; font-weight: bold; margin:0;}
    .header-box p { font-size: 1.1rem; margin-top: 5px; opacity: 0.9;}
    </style>
    <div class="header-box">
        <h1>EDUPLAN IA 🇵🇪</h1>
        <p>Sistema Experto de Planificación Curricular - MINEDU</p>
    </div>
""", unsafe_allow_html=True)

# Estado de sesión para almacenar resultados
if 'resultado_generado' not in st.session_state:
    st.session_state.resultado_generado = None
if 'tipo_doc_actual' not in st.session_state:
    st.session_state.tipo_doc_actual = None

with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/1/14/Logo_del_Ministerio_de_Educaci%C3%B3n_del_Per%C3%BA.svg/1024px-Logo_del_Ministerio_de_Educaci%C3%B3n_del_Per%C3%BA.svg.png", width=200)
    st.title("⚙️ Configuración")
    api_key = st.text_input("API Key (ZhipuAI / OpenAI compatible):", type="password", help="Ingresa tu clave de API para generar los documentos.")
    st.markdown("---")
    st.markdown("**Desarrollado para docentes de la EBR.**")
    st.markdown("*Genera documentos listos para impresión, alineados a rúbricas de evaluación del Minedu.*")

tab1, tab2, tab3 = st.tabs(["📝 Sesión de Aprendizaje", "📚 Unidad Didáctica", "📅 Programación Anual"])

def limpiar_dependencias(key_grado, key_area):
    """Limpia la memoria caché de grado y área para evitar que los selectores se congelen."""
    if key_grado in st.session_state:
        del st.session_state[key_grado]
    if key_area in st.session_state:
        del st.session_state[key_area]

def form_ui(tipo_doc):
    # SOLUCIÓN: Quitamos 'with st.form' y agregamos 'key' a cada widget para que se actualicen en tiempo real
    col1, col2 = st.columns(2)
    
    grado_key = f"gra_{tipo_doc}"
    area_key = f"area_{tipo_doc}"
    
    with col1:
        docente = st.text_input("Nombre del Docente", help="Aparecerá en los datos informativos.", key=f"doc_{tipo_doc}")
        ie = st.text_input("Institución Educativa", help="Nombre o número del colegio.", key=f"ie_{tipo_doc}")
        
        # Al seleccionar el nivel, el script se actualiza al instante y resetea las memorias
        nivel = st.selectbox(
            "Nivel Educativo", 
            list(NIVELES_GRADOS.keys()), 
            key=f"niv_{tipo_doc}",
            on_change=limpiar_dependencias,
            args=(grado_key, area_key)
        )
        
        # Y muestra los grados correspondientes al nivel de arriba
        grado = st.selectbox("Grado/Edad", NIVELES_GRADOS[nivel], key=grado_key)
        
    with col2:
        # Muestra las áreas del nivel seleccionado
        area = st.selectbox("Área Curricular", AREAS_NIVEL[nivel], key=area_key)
        tema = st.text_input("Tema / Título Principal", help="Ej. 'Conocemos el ciclo del agua'", key=f"tema_{tipo_doc}")
        enfoque = st.selectbox("Enfoque Transversal", ENFOQUES_TRANSVERSALES, key=f"enf_{tipo_doc}")
        
        if tipo_doc == "Sesión de Aprendizaje":
            minutos = st.number_input("Duración (minutos)", min_value=15, max_value=300, value=90, step=15, help="Escoge el tiempo en minutos para la sesión.", key=f"dur_{tipo_doc}")
            duracion = f"{minutos} minutos"
        else:
            duracion = st.text_input("Duración", value="Aproximadamente 4 semanas", key=f"dur_{tipo_doc}")

        instrumento = None
        if tipo_doc == "Sesión de Aprendizaje":
            instrumento = st.selectbox(
                "📊 Instrumento de Evaluación", 
                INSTRUMENTOS_EVALUACION, 
                help="Selecciona el instrumento según CNEB que la IA diseñará al final de la sesión.",
                key=f"inst_{tipo_doc}"
            )

    contexto = st.text_area("Situación Significativa / Contexto local", help="Describe brevemente la realidad de los estudiantes o problemática local. (Opcional pero recomendado).", key=f"ctx_{tipo_doc}")
    
    producto = ""
    if tipo_doc == "Unidad Didáctica":
        producto = st.text_input("Producto Final Esperado", help="Ej. 'Un afiche sobre el cuidado del medio ambiente'", key=f"prod_{tipo_doc}")

    st.markdown("<br>", unsafe_allow_html=True)
    # Reemplazamos st.form_submit_button por st.button
    submitted = st.button(f"🚀 Generar {tipo_doc} con IA", key=f"btn_{tipo_doc}", use_container_width=True)
    
    if submitted:
        if not api_key:
            st.error("⚠️ Por favor, ingresa tu API Key en la barra lateral izquierda.")
            return None
        if not docente or not tema:
            st.warning("⚠️ El Nombre del docente y el Tema son obligatorios.")
            return None
            
        return {
            "tipo_doc": tipo_doc,
            "docente": docente, "ie": ie, "nivel": nivel, "grado": grado,
            "area": area, "tema": tema, "enfoque": enfoque, "duracion": duracion,
            "contexto": contexto, "producto": producto, "instrumento": instrumento
        }
    return None

# Renderizado de Formularios
datos_generar = None

with tab1:
    st.subheader("Planificar Sesión de Aprendizaje")
    st.info("Incluye Momentos didácticos, Instrumento de Evaluación (Rúbrica/Lista) y Ficha de trabajo.")
    datos = form_ui("Sesión de Aprendizaje")
    if datos: datos_generar = datos

with tab2:
    st.subheader("Planificar Unidad Didáctica / Experiencia")
    st.info("Incluye Situación Significativa, Propósitos y Secuencia de Actividades.")
    datos = form_ui("Unidad Didáctica")
    if datos: datos_generar = datos

with tab3:
    st.subheader("Programación Anual")
    st.info("Incluye Matriz anualizada, Calendarización y Organización de Unidades.")
    datos = form_ui("Programación Anual")
    if datos: datos_generar = datos

# --- 5. EJECUCIÓN Y LLAMADA A LA IA ---
if datos_generar:
    with st.spinner(f"🧠 La IA está analizando el CNEB y redactando la {datos_generar['tipo_doc']}... Esto puede tomar unos 30 segundos."):
        try:
            client = ZhipuAI(api_key=api_key)
            prompt = generar_prompt_cneb(datos_generar['tipo_doc'], datos_generar)
            
            response = client.chat.completions.create(
                model="glm-4",  # Asegúrate de usar el modelo correcto según tu tier (glm-3-turbo o glm-4)
                messages=[
                    {"role": "system", "content": "Eres un redactor técnico experto en educación. Usa Markdown, formatea tablas correctamente y sé detallado. Nunca dejes celdas vacías en las tablas."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            )
            
            # Limpiamos el texto por si la IA añade bloques de código residuales al inicio/fin
            texto_respuesta = response.choices[0].message.content
            texto_respuesta = re.sub(r'^```markdown\n?', '', texto_respuesta)
            texto_respuesta = re.sub(r'\n?```$', '', texto_respuesta)
            
            st.session_state.resultado_generado = texto_respuesta
            st.session_state.tipo_doc_actual = datos_generar['tipo_doc']
            st.success("¡Documento pedagógico generado exitosamente!")
            
        except Exception as e:
            error_msg = str(e)
            if "authentication" in error_msg.lower() or "api_key" in error_msg.lower() or "401" in error_msg:
                st.error("❌ Error de Autorización: Tu API Key de ZhipuAI es incorrecta o inválida.")
            else:
                st.error(f"❌ Ocurrió un error al conectar con la IA: {error_msg}")

# --- 6. PREVISUALIZACIÓN Y DESCARGA ---
if st.session_state.resultado_generado:
    st.markdown("---")
    st.subheader(f"📄 Previsualización: {st.session_state.tipo_doc_actual}")
    
    with st.expander("Ver documento redactado (Markdown)", expanded=True):
        st.markdown(st.session_state.resultado_generado)

    # Generación y Descarga de DOCX
    st.markdown("### 📥 Exportar Documento")
    with st.spinner("Empaquetando en formato Word Profesional..."):
        try:
            docx_file = markdown_to_docx(st.session_state.resultado_generado)
            file_name = f"{st.session_state.tipo_doc_actual.replace(' ', '_')}_{int(time.time())}.docx"
            
            st.download_button(
                label="⬇️ Descargar Documento en Word (.docx)",
                data=docx_file,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            st.caption("El documento incluye formatos de tabla, márgenes oficiales, encabezados y pies de página listos para imprimir.")
        except Exception as e:
            st.error(f"Error al generar el archivo Word: {str(e)}")
